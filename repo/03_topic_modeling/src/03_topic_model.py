import pandas as pd
import csv
from bertopic import BERTopic
from bertopic.representation import MaximalMarginalRelevance
from sentence_transformers import SentenceTransformer
from sklearn.feature_extraction.text import CountVectorizer
from docx import Document
from openai import OpenAI

# load sermon texts
df = pd.read_csv("input/sermons.csv")
# only keep sermons within 2 year window to an election
df = df.dropna(subset= ['weeks_to_nearest_election'])
df = df[(df['weeks_to_nearest_election'] >= -104) & (df['weeks_to_nearest_election'] <= 104)]
# df where all sermons are within 4 week window to an election tuesday
df_tuesday = df.copy().dropna(subset= ['weeks_to_nearest_tuesday'])
df_tuesday = df_tuesday[(df_tuesday['weeks_to_nearest_tuesday'] >= -4) & (df_tuesday['weeks_to_nearest_tuesday'] <= 4)]


def run_all(df, docs_type):
    def run_topic_model(docs, n, m):
        # define words not to be in topics (candidate names and basic words)
        candidate_keywords = [
            'Bill Clinton', 'Hillary Clinton', 'Clinton', 
            'George W. Bush', 'George Bush', 'President Bush', 'Mr. Bush', 
            'Barack Obama', 'Obama', 
            'Joe Biden', 'Biden', 
            'John McCain', 'McCain', 
            'Mitt Romney', 'Romney', 
            'Donald Trump', 'Trump', 
            'Bernie Sanders', 'Kamala Harris', 'Al Gore'
        ]
        candidate_keywords_ext = [
            'Bill', 'Hillary', 'John', 'Barack', 'Joe', 'George', 'Mitt', 'Donald', 
            'Bernie', 'Sanders', 'Kamala', 'Harris', 'Al', 'Gore', 
            'Jimmy Carter', 'Carter', 'Jimmy', 'Richard Nixon', 'Nixon', 'Richard', 
            'John F. Kennedy', 'John F Kennedy', 'John Kennedy', 'Kennedy', 'John', 
            'Abraham Lincoln', 'Lincoln', 'Abraham', 
            'Hussein', 'Bush', 'Trumpet'
        ]
        candidate_keywords_total = candidate_keywords + candidate_keywords_ext
        stop_words = candidate_keywords_total + [name.lower() for name in candidate_keywords_total] + [
            "a", "an", "the", "this", "that", "these", "those",
            "and", "or", "but", "if", "then", "not", "for",
            "of", "by", "at", "in", "on", "there", 
            "to", "be", "is", "are", "was", "were", "will",
            "do", "did", "have", "has", "had",
            "I", "you", "he", "she", "it", "we", "they",
            "me", "him", "her", "us", "them",
            "my", "your", "yours", "his", "her", "hers", 
            "its", "our", "ours", "their", "theirs",
            "one", "some", "all", "none", "no", 
            "who", "what", "when", "where", "why", "how",
            "as", "with", "so", "about", "because"
        ]

        # initialize embedding model to capture word meaning, not word frequency
        embedding_model = SentenceTransformer("all-MiniLM-L6-v2")
        # ignore words from list of stop words, and words in less than 5 docs or more than 90% of docs
        vectorizer_model = CountVectorizer(stop_words= stop_words, min_df= 5, max_df= 0.8)
        # initialize representation model to reduce similarity between topic representative words
        representation_model = MaximalMarginalRelevance(diversity= 0.4)

        # load and fit BERTopic with n initial topics
        topic_model = BERTopic(
            embedding_model= embedding_model, 
            vectorizer_model= vectorizer_model, 
            representation_model= representation_model, 
            nr_topics= n
        )
        topics, probs = topic_model.fit_transform(docs)
        
        # reduce topics
        topic_model.reduce_topics(docs, nr_topics= m+1)
        topics_reduced, probs_reduced = topic_model.transform(docs)

        # try to assign each outlier to best matching topic
        try:
            topics_reduced = topic_model.reduce_outliers(docs, topics_reduced, probabilities= probs_reduced, strategy= 'probabilities')
        except ValueError as err:
            if str(err) != "No outliers to reduce.":
                raise err

        # get results
        df_topic_model_doc_info = topic_model.get_document_info(docs)
        # group entries by topic then sort on probabilities
        df_topic_model_doc_info = df_topic_model_doc_info.sort_values(by= ["Topic", "Probability"], ascending= [True, False])
        # order columns
        df_topic_model_doc_info = df_topic_model_doc_info[["Topic", "Name", "Representation", "Document"]]

        return topic_model, df_topic_model_doc_info
    
    def write_document_output(df_topic_model_doc_info, docs_type):
        # generate a topic label with gpt
        def get_topic_label(name, df_topic):
            topic_label_context = []
            topic_label_context.append(f"Name: {name}")
            topic_label_context.append(f"Top 10 words: {', '.join(df_topic['Representation'].iloc[0])}")
            topic_label_context.append("Example sentences:")

            for j, row in enumerate(df_topic.itertuples(index= False), start= 1):
                topic_label_context.append(f"{j}. {row.Document}.")

            topic_label_context = "\n".join(topic_label_context)

            
            client = OpenAI()
            prompt = (
                "You are an expert at labeling topics generated by topic models. "
                "Given a topic's name, top representative words, and example sentences, "
                "provide a short, descriptive topic label that summarizes the topic. "
                "Only return the label, no explanation."
            )
            response = client.chat.completions.create(
                model= "gpt-4o",
                messages= [
                    {"role": "system", "content": prompt},
                    {"role": "user", "content": topic_label_context}
                ],
                temperature= 0.3
            )

            return response.choices[0].message.content
        
        topic_label_map = dict()

        # keep only the first entries for each topic
        df_for_document = df_topic_model_doc_info.groupby("Topic").head()
        topics = df_for_document["Name"].unique()

        doc = Document()
        doc.add_heading("Topic Model Output for Sermons Given within 4 Weeks of the Nearest Election", level= 1)
        doc.add_page_break()

        for i in range(len(topics)):
            df_topic = df_for_document[df_for_document["Name"] == topics[i]]
            topic_name = topics[i]
            topic_label = get_topic_label(topic_name, df_topic)
            #topic_label = topic_name
            topic_label_map[i-1] = topic_label

            if i == 0:
                doc.add_paragraph(f"Outlier Topic: {topic_label}", style= 'Heading 2')
            else:
                doc.add_paragraph(f"Topic {i-1}: {topic_label}", style= 'Heading 2')
            doc.add_paragraph(f"Top 10 words: {', '.join(df_topic['Representation'].iloc[0])}")
            doc.add_paragraph("Example sentences:")

            for j, row in enumerate(df_topic.itertuples(index= False), start= 1):
                doc.add_paragraph(f"{j}. {row.Document}.")

            if i != len(topics) - 1:
                doc.add_page_break()

        doc.save(f"output/{docs_type}/topic_model_output.docx")

        return topic_label_map

    def get_regression_data(topic_model, topic_label_map, df, df_with_docs, docs, regression):
        # transform docs into existing topics
        topics_fit, probs_fit = topic_model.transform(docs)

        df_with_topics = df_with_docs.copy()
        df_with_topics['topic'] = topics_fit
        df_with_topics['prob'] = probs_fit
        df_with_topics['label'] = df_with_topics['topic'].map(topic_label_map)

        # get topic labels
        df_topic_labels = (
            df_with_topics[['topic', 'label']]
            .drop_duplicates()
            .assign(topic= lambda d : d['topic'].astype(int))
            .sort_values(by= 'topic')
        )
        df_topic_labels.to_csv(f"intermediate/regression/{regression}/labels.csv", index=False)

        # for all sermon texts, track topics that are mentioned within the text
        mentioned_topics = (
            df_with_topics[df_with_topics['prob'] > 0.5]
            .groupby('sermontext')['topic']
            .apply(lambda x: list(set(x)))
            .reset_index()
            .rename(columns= {'topic': 'mentioned_topics'})
        )

        df = df.merge(mentioned_topics, on= 'sermontext', how= 'left')
        df['mentioned_topics'] = df['mentioned_topics'].apply(
            lambda x : x if isinstance(x, list) else []
        )
        df['mentioned_topic_labels'] = df['mentioned_topics'].apply(
            lambda topics : [topic_label_map[topic] for topic in topics]
        )

        # drop edge case if in_window_dummy regression
        if regression == 'pattern_matching_in_window':
            df = df[(df['weeks_to_nearest_tuesday'] > -26) & (df['weeks_to_nearest_tuesday'] < 26)]

        df = df[['index', 'date', 'weeks_to_nearest_tuesday', 'weeks_to_nearest_election', 'mentioned_topics', 'mentioned_topic_labels', 'sermontext']]
        df.to_csv(f"intermediate/regression/{regression}/data.csv", index= False)

    # get docs for topic modeling
    if docs_type == 'pattern_matching':
        df_with_docs = pd.read_csv(f"intermediate/{docs_type}/all_docs.csv")
        df_tuesday_with_docs = pd.read_csv(f"intermediate/{docs_type}/tuesday_docs.csv")
        docs_tuesday = df_tuesday_with_docs['doc'].tolist()
    elif docs_type == 'gpt':
        df_with_docs = pd.read_csv(f"intermediate/{docs_type}/tuesday_docs.csv")
    df_election_with_docs = pd.read_csv(f"intermediate/{docs_type}/election_docs.csv")

    docs = df_with_docs['doc'].tolist()
    docs_election = df_election_with_docs['doc'].tolist()

    # run topic model on election docs
    topic_model, df_topic_model_doc_info = run_topic_model(docs_election, 64, 16)
    # save results
    df_topic_model_doc_info = df_topic_model_doc_info[["Topic", "Name", "Representation", "Document"]]
    df_topic_model_doc_info.to_csv(f"output/{docs_type}/topic_model_output.csv", index= False, quoting= csv.QUOTE_ALL)

    # write document output and get topic label mapping
    topic_label_map = write_document_output(df_topic_model_doc_info, docs_type)

    # get regression data
    if docs_type == 'pattern_matching':
        get_regression_data(topic_model, topic_label_map, df, df_with_docs, docs, 'pattern_matching_in_window')
        df_tuesday = df.copy().dropna(subset= ['weeks_to_nearest_tuesday'])
        df_tuesday = df_tuesday[(df_tuesday['weeks_to_nearest_tuesday'] >= -4) & (df_tuesday['weeks_to_nearest_tuesday'] <= 4)]
        get_regression_data(topic_model, topic_label_map, df_tuesday, df_tuesday_with_docs, docs_tuesday, 'pattern_matching_pre_tuesday')
    elif docs_type == 'gpt':
        get_regression_data(topic_model, topic_label_map, df, df_with_docs, docs, 'gpt_pre_tuesday')

run_all(df.copy(), 'pattern_matching')
run_all(df_tuesday.copy(), 'gpt')
